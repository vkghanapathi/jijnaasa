#!/usr/bin/env python3
"""
Fix bs-ns12.md:
1. Remove added N.M sentence numbering from commentary lines
2. Add sentence count to each episode heading
3. Generate TOC at top
4. Output both .md and .docx
"""

import re
from pathlib import Path

INPUT = Path('/Users/dharmaposhanam/Documents/GitHub/Jijnaasa/data/dvaita/bs-ns12.md')
OUT_MD = INPUT
OUT_DOCX = INPUT.with_suffix('.docx')

# Pattern matching the added numbering: "N.M " at start of line (e.g., "1.1 ", "23.45 ")
NUMBERING_RE = re.compile(r'^(\d+)\.(\d+)\s')

# Episode heading: "# N. heading_text"
EPISODE_RE = re.compile(r'^#\s+(\d+)\.\s+(.+)$')

# Section heading: "# भाग N — ..."
SECTION_RE = re.compile(r'^#\s+भाग\s')

# Sub-heading: "## न्यायसुधा"
SUBHEAD_RE = re.compile(r'^##\s')

# Samiksha / AV table lines
TABLE_RE = re.compile(r'^\|')

# Main title line
TITLE_RE = re.compile(r'^#\s+न्यायसुधा\s+—')


def parse_episodes(lines):
    """Parse the file and collect episodes with their sentence counts."""
    episodes = []
    current_episode = None
    current_section = None
    sentence_count = 0

    for line in lines:
        s = line.rstrip()

        # Track section
        if SECTION_RE.match(s):
            current_section = s
            continue

        # Episode heading
        m = EPISODE_RE.match(s)
        if m:
            # Save previous episode
            if current_episode is not None:
                current_episode['sentences'] = sentence_count
            num = int(m.group(1))
            heading = m.group(2)
            current_episode = {
                'num': num,
                'heading': heading,
                'section': current_section,
            }
            episodes.append(current_episode)
            sentence_count = 0
            continue

        # Count commentary sentences (lines starting with N.M)
        if NUMBERING_RE.match(s):
            sentence_count += 1

    # Save last episode
    if current_episode is not None:
        current_episode['sentences'] = sentence_count

    return episodes


def fix_file(lines, episodes):
    """Remove numbering, add sentence counts to headings."""
    # Build episode lookup: (num, heading) -> sentence_count
    ep_lookup = {}
    for ep in episodes:
        ep_lookup[(ep['num'], ep['heading'])] = ep['sentences']

    out = []
    for line in lines:
        s = line.rstrip()

        # Fix episode heading: add sentence count
        m = EPISODE_RE.match(s)
        if m:
            num = int(m.group(1))
            heading = m.group(2)
            count = ep_lookup.get((num, heading), 0)
            if count > 0:
                out.append(f'# {num}. {heading} ({count} sentences)')
            else:
                out.append(s)
            continue

        # Remove added numbering from commentary lines
        m2 = NUMBERING_RE.match(s)
        if m2:
            # Strip the "N.M " prefix, keep rest
            text = NUMBERING_RE.sub('', s)
            out.append(text)
            continue

        out.append(s)

    return out


def generate_toc(episodes):
    """Generate table of contents."""
    toc = []
    toc.append('# विषयानुक्रमणिका (Table of Contents)')
    toc.append('')

    current_section = None
    for ep in episodes:
        if ep['section'] != current_section:
            current_section = ep['section']
            if current_section:
                toc.append(f'## {current_section.lstrip("# ")}')
                toc.append('')

        count_str = f' — {ep["sentences"]} sentences' if ep['sentences'] > 0 else ''
        toc.append(f'{ep["num"]}. {ep["heading"]}{count_str}')

    toc.append('')
    toc.append('---')
    toc.append('')
    return toc


def write_docx(lines, episodes, outpath):
    """Write docx output."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(12)

    for line in lines:
        s = line.rstrip()

        if not s:
            continue

        # Main title
        if TITLE_RE.match(s):
            p = doc.add_heading(s.lstrip('# '), level=0)
            continue

        # Section heading (भाग)
        if SECTION_RE.match(s):
            doc.add_heading(s.lstrip('# '), level=1)
            continue

        # Episode heading
        m = EPISODE_RE.match(s)
        if m:
            doc.add_heading(s.lstrip('# '), level=2)
            continue

        # Sub-heading (न्यायसुधा)
        if SUBHEAD_RE.match(s):
            doc.add_heading(s.lstrip('# '), level=3)
            continue

        # Separator
        if s == '---':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('─' * 40)
            run.font.color.rgb = RGBColor(180, 180, 180)
            continue

        # Table lines (AV box, Samiksha box)
        if TABLE_RE.match(s):
            # Extract text between pipes
            text = s.strip('| ').replace('\\|', '|')
            text = text.replace('&nbsp;', '   ')
            if '**' in text:
                text = text.replace('**', '')
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
            elif text.strip() == ':---':
                continue
            else:
                doc.add_paragraph(text)
            continue

        # Metadata lines
        if s.startswith('**') and '—' in s:
            text = s.replace('**', '')
            p = doc.add_paragraph()
            parts = text.split('—', 1)
            run = p.add_run(parts[0].strip())
            run.bold = True
            if len(parts) > 1:
                p.add_run(f' — {parts[1].strip()}')
            continue

        # Regular text
        doc.add_paragraph(s)

    doc.save(str(outpath))


def main():
    lines = INPUT.read_text(encoding='utf-8').splitlines()

    # Parse episodes
    episodes = parse_episodes(lines)

    print(f'Found {len(episodes)} episodes')
    total_sentences = sum(ep['sentences'] for ep in episodes)
    print(f'Total sentences: {total_sentences}')
    print()

    # Fix numbering and add counts
    fixed = fix_file(lines, episodes)

    # Generate TOC
    toc = generate_toc(episodes)

    # Find insertion point for TOC (after the header block, before first भाग)
    insert_idx = 0
    for i, line in enumerate(fixed):
        if SECTION_RE.match(line):
            insert_idx = i
            break

    # Insert TOC
    final = fixed[:insert_idx] + toc + fixed[insert_idx:]

    # Write markdown
    OUT_MD.write_text('\n'.join(final) + '\n', encoding='utf-8')
    print(f'Written fixed markdown: {OUT_MD}')

    # Write docx
    write_docx(final, episodes, OUT_DOCX)
    print(f'Written docx: {OUT_DOCX}')


if __name__ == '__main__':
    main()

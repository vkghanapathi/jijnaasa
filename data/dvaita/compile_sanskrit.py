#!/usr/bin/env python3
"""
compile_sanskrit.py
-------------------
Reads 4 Sanskrit translation markdown files and compiles them into:
  1. 05-ns-review.docx  (formatted Word document in Devanagari)
  2. 05-ns-review.md    (combined Markdown reference)
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ─── Configuration ───────────────────────────────────────────────────────────

BASE_DIR = Path("/Users/dharmaposhanam/Documents/GitHub/Jijnaasa/data/dvaita")

SOURCE_FILES = [
    BASE_DIR / "sanskrit_eps_01_22.md",
    BASE_DIR / "sanskrit_eps_23_44.md",
    BASE_DIR / "sanskrit_eps_45_66.md",
    BASE_DIR / "sanskrit_eps_67_end.md",
]

OUT_DOCX = BASE_DIR / "05-ns-review.docx"
OUT_MD = BASE_DIR / "05-ns-review.md"

# ─── Parsing ─────────────────────────────────────────────────────────────────

# Episode heading: ### प्रकरणम् N: Title
EPISODE_RE = re.compile(r'^###\s+प्रकरणम्\s+(\d+)\s*[:\.\—–-]\s*(.+)$')

# Part heading: ## भागः ...
PART_RE = re.compile(r'^##\s+(.+)$')

# Section heading within episode: **पूर्वपक्षः** etc.
SECTION_BOLD_RE = re.compile(r'^\*\*(.+?)\*\*\s*$')

# Top-level heading
TITLE_RE = re.compile(r'^#\s+(.+)$')

# Horizontal rule
HR_RE = re.compile(r'^---+$')


def parse_markdown(filepath: Path):
    """Parse a Sanskrit translation markdown file into structured data."""
    lines = filepath.read_text(encoding='utf-8').splitlines()

    episodes = []
    current_part = None
    current_episode = None
    current_section = None

    for line in lines:
        s = line.rstrip()

        # Skip title lines (# ...)
        m = TITLE_RE.match(s)
        if m and not PART_RE.match(s) and not EPISODE_RE.match(s):
            continue

        # Part heading
        m = PART_RE.match(s)
        if m:
            current_part = m.group(1).strip()
            continue

        # Episode heading
        m = EPISODE_RE.match(s)
        if m:
            if current_episode:
                episodes.append(current_episode)
            num = int(m.group(1))
            title = m.group(2).strip()
            current_episode = {
                'num': num,
                'title': title,
                'part': current_part,
                'sections': [],
                'current_section_lines': [],
                'current_section_name': None,
            }
            current_section = None
            continue

        if current_episode is None:
            continue

        # Horizontal rule — skip
        if HR_RE.match(s):
            continue

        # Section bold heading (**पूर्वपक्षः** etc.)
        m = SECTION_BOLD_RE.match(s)
        if m:
            # Save previous section
            if current_episode['current_section_name']:
                current_episode['sections'].append({
                    'name': current_episode['current_section_name'],
                    'text': '\n'.join(current_episode['current_section_lines']).strip(),
                })
            current_episode['current_section_name'] = m.group(1).strip()
            current_episode['current_section_lines'] = []
            continue

        # Regular line — accumulate
        if current_episode['current_section_name']:
            current_episode['current_section_lines'].append(s)
        elif s.strip():
            # Lines before first section — treat as intro
            if not current_episode['sections'] and not current_episode.get('intro'):
                current_episode['intro'] = []
            if 'intro' in current_episode:
                current_episode['intro'].append(s)

    # Save last episode
    if current_episode:
        if current_episode['current_section_name']:
            current_episode['sections'].append({
                'name': current_episode['current_section_name'],
                'text': '\n'.join(current_episode['current_section_lines']).strip(),
            })
        episodes.append(current_episode)

    # Clean up temp fields
    for ep in episodes:
        ep.pop('current_section_lines', None)
        ep.pop('current_section_name', None)

    return episodes


# ─── DOCX Building ──────────────────────────────────────────────────────────

def setup_styles(doc: Document):
    """Set up styles for Sanskrit document."""
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Cambria"
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2

    title = styles["Title"]
    title.font.name = "Cambria"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x1A, 0x1A, 0x5C)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Cambria"
    subtitle.font.size = Pt(13)
    subtitle.font.italic = True
    subtitle.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(20)

    h1 = styles["Heading 1"]
    h1.font.name = "Cambria"
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)

    h2 = styles["Heading 2"]
    h2.font.name = "Cambria"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x1A, 0x1A, 0x5C)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)

    h3 = styles["Heading 3"]
    h3.font.name = "Cambria"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)


def add_rich_paragraph(doc, text):
    """Add paragraph with inline bold (**...**) support."""
    text = text.strip()
    if not text:
        return
    para = doc.add_paragraph()
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        if i % 2 == 1:
            run.bold = True


def build_docx(all_episodes):
    """Build the DOCX from parsed episodes."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.0)

    setup_styles(doc)

    # Title
    doc.add_paragraph("न्यायसुधा — संस्कृतसमीक्षा", style="Title")
    doc.add_paragraph(
        "ब्रह्मसूत्र १.१.१–१.१.२ विषये न्यायसुधायाः पूर्वपक्ष-सिद्धान्त-विश्लेषणम्",
        style="Subtitle"
    )

    # Author info
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_para.add_run(
        "जिज्ञासा-अधिकरण-तुलनात्मक-विश्लेषण-प्रकल्पः\n"
        "अन्वेषकः: एन्. कुवलयदत्तः\n"
        "मार्गदर्शकः: प्रो. एम्.एल्.एन्. भट्टः (राष्ट्रियसंस्कृतविश्वविद्यालयः, तिरुपतिः)\n"
        "परामर्शदाता: डॉ. वंशीकृष्ण-घनपाठी"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    info_para.paragraph_format.space_after = Pt(30)

    doc.add_paragraph("─" * 60)

    # Track parts
    current_part = None

    for ep in all_episodes:
        part = ep.get('part', '')
        if part and part != current_part:
            current_part = part
            if 'जन्मा' in part or 'Janm' in part:
                doc.add_page_break()
            doc.add_heading(part, level=1)

        # Episode heading
        doc.add_heading(f"प्रकरणम् {ep['num']}: {ep['title']}", level=2)

        # Intro if any
        if ep.get('intro'):
            for line in ep['intro']:
                if line.strip():
                    add_rich_paragraph(doc, line)

        # Sections
        for sec in ep['sections']:
            doc.add_heading(sec['name'], level=3)
            # Split by double newlines for paragraphs
            paras = sec['text'].split('\n\n')
            for para_text in paras:
                para_text = para_text.strip()
                if not para_text:
                    continue
                # Join single newlines within paragraph
                para_text = ' '.join(para_text.split('\n'))
                add_rich_paragraph(doc, para_text)

    doc.save(str(OUT_DOCX))
    print(f"  Saved: {OUT_DOCX}")


# ─── Markdown Compilation ───────────────────────────────────────────────────

def build_markdown(all_episodes):
    """Combine all episodes into a single markdown file."""
    md = []
    md.append("# न्यायसुधा — संस्कृतसमीक्षा\n")
    md.append("**ब्रह्मसूत्र १.१.१–१.१.२ विषये न्यायसुधायाः पूर्वपक्ष-सिद्धान्त-विश्लेषणम्**\n")
    md.append("---\n")

    current_part = None
    for ep in all_episodes:
        part = ep.get('part', '')
        if part and part != current_part:
            current_part = part
            md.append(f"\n## {part}\n")

        md.append(f"\n### प्रकरणम् {ep['num']}: {ep['title']}\n")

        if ep.get('intro'):
            for line in ep['intro']:
                md.append(line)
            md.append("")

        for sec in ep['sections']:
            md.append(f"\n**{sec['name']}**\n")
            md.append(sec['text'])
            md.append("")

        md.append("\n---\n")

    content = '\n'.join(md)
    OUT_MD.write_text(content, encoding='utf-8')
    print(f"  Saved: {OUT_MD}")


# ─── Main ────────────────────────────────────────────────────────────────────

def main():
    print("Compiling Sanskrit translation files...")

    all_episodes = []
    for filepath in SOURCE_FILES:
        if not filepath.exists():
            print(f"  WARNING: {filepath.name} not found")
            continue
        episodes = parse_markdown(filepath)
        print(f"  {filepath.name}: {len(episodes)} episodes")
        all_episodes.extend(episodes)

    print(f"\nTotal episodes: {len(all_episodes)}")

    print("\nBuilding DOCX...")
    build_docx(all_episodes)

    print("Building Markdown...")
    build_markdown(all_episodes)

    print(f"\nDone. {len(all_episodes)} episodes compiled.")


if __name__ == "__main__":
    main()

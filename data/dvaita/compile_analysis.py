#!/usr/bin/env python3
"""
compile_analysis.py
-------------------
Reads 4 agent-output files containing episode-wise analysis of Nyāya Sudhā,
extracts the analytical notes, and compiles them into:
  1. ns-analysis.docx  (formatted Word document)
  2. ns-analysis.md    (Markdown reference)

Source files (JSONL agent logs):
  - Jijñāsādhikaraṇa Episodes 1-31   (BS 1.1.1)
  - Jijñāsādhikaraṇa Episodes 32-57  (BS 1.1.1)
  - Jijñāsādhikaraṇa Episodes 58-83  (BS 1.1.1)
  - Janmādhikaraṇa Episodes 1-6       (BS 1.1.2)
"""

import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# ─── Configuration ───────────────────────────────────────────────────────────

TASK_DIR = Path("/private/tmp/claude-501/-Users-dharmaposhanam-Documents-GitHub-Jijnaasa/tasks")
OUTPUT_DIR = Path("/Users/dharmaposhanam/Documents/GitHub/Jijnaasa/data/dvaita")

# Map: (file_id, part_label, ordering_key)
# Ordered so that Jijñāsādhikaraṇa comes first, then Janmādhikaraṇa
SOURCE_FILES = [
    ("a8d5f7c7105220708.output", "jijnasa", 1),   # Eps 1-31
    ("aca7595be6e756e58.output", "jijnasa", 2),    # Eps 32-57
    ("a91783d14580c63c5.output", "jijnasa", 3),    # Eps 58-83
    ("a62b1f5be63752075.output", "janma",   4),    # Janma Eps 1-6
]


# ─── Extraction ──────────────────────────────────────────────────────────────

def extract_analysis_text(filepath: Path) -> str:
    """Extract the final assistant text block containing '### Episode' from a JSONL log."""
    candidate = ""
    with open(filepath, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            try:
                data = json.loads(line)
            except json.JSONDecodeError:
                continue
            if data.get("type") != "assistant":
                continue
            content = data.get("message", {}).get("content", "")
            if isinstance(content, list):
                for item in content:
                    if isinstance(item, dict) and item.get("type") == "text":
                        text = item["text"]
                        if "### Episode" in text and len(text) > 500:
                            candidate = text  # keep the last (most complete) one
            elif isinstance(content, str) and "### Episode" in content and len(content) > 500:
                candidate = content
    return candidate


def parse_episodes(text: str):
    """
    Parse episode blocks and summary from extracted text.
    Returns: (episodes_list, summary_text)
    Each episode: dict with keys 'heading', 'body'
    """
    # Strip preamble (before first ### Episode)
    match = re.search(r'^### Episode', text, re.MULTILINE)
    if match:
        text = text[match.start():]

    # Split into episodes + trailing summary
    # Pattern: ### Episode ... (everything until next ### Episode or ## Summary)
    parts = re.split(r'^(?=### Episode |## Summary)', text, flags=re.MULTILINE)

    episodes = []
    summary = ""
    for part in parts:
        part = part.strip()
        if not part:
            continue
        if part.startswith("## Summary"):
            summary = part
        elif part.startswith("### Episode"):
            # Extract heading line
            lines = part.split("\n", 1)
            heading = lines[0].lstrip("#").strip()
            body = lines[1].strip() if len(lines) > 1 else ""
            # Strip trailing horizontal rules from body
            body = re.sub(r'\n---+\s*$', '', body).strip()
            episodes.append({"heading": heading, "body": body})

    return episodes, summary


# ─── DOCX Formatting ────────────────────────────────────────────────────────

def setup_styles(doc: Document):
    """Set up custom styles for the document."""
    styles = doc.styles

    # Modify existing Normal style
    normal = styles["Normal"]
    normal.font.name = "Cambria"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    # Title style
    title_style = styles["Title"]
    title_style.font.name = "Cambria"
    title_style.font.size = Pt(22)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x5C)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(4)

    # Subtitle style
    subtitle = styles["Subtitle"]
    subtitle.font.name = "Cambria"
    subtitle.font.size = Pt(13)
    subtitle.font.italic = True
    subtitle.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(20)

    # Heading 1 — Part headers
    h1 = styles["Heading 1"]
    h1.font.name = "Cambria"
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)  # dark red
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)

    # Heading 2 — Episode headings
    h2 = styles["Heading 2"]
    h2.font.name = "Cambria"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x1A, 0x1A, 0x5C)  # dark blue
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)

    # Heading 3 — Sub-sections
    h3 = styles["Heading 3"]
    h3.font.name = "Cambria"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)

    # Create a custom style for summary sections
    try:
        summary_style = styles.add_style("SummaryHeading", WD_STYLE_TYPE.PARAGRAPH)
        summary_style.base_style = styles["Heading 2"]
        summary_style.font.color.rgb = RGBColor(0x00, 0x66, 0x00)  # dark green
    except ValueError:
        pass  # style already exists


def add_body_paragraphs(doc: Document, body_text: str):
    """
    Parse body text containing bold markers (**...**) and horizontal rules (---),
    and add formatted paragraphs to the document.
    """
    # Split on horizontal rules
    sections = re.split(r'\n---+\n', body_text)

    for section in sections:
        section = section.strip()
        if not section:
            continue

        # Split into logical paragraphs (double newline or bold-start patterns)
        # Each sub-section may start with **Bold Label:**
        paragraphs = re.split(r'\n(?=\*\*)', section)

        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue

            # Check if this is a sub-heading like **Pūrvapakṣa (Advaita position targeted):**
            sub_heading_match = re.match(
                r'\*\*(Pūrvapakṣa.*?|Jayatīrtha\'s argument.*?|Advaita counter.*?|'
                r'Dvaita presuppositions.*?|Summary.*?|Key.*?|Thesis.*?|Structural.*?|'
                r'Methodological.*?|Central.*?|Conclusion.*?)\*\*[:\s]*(.*)',
                para_text, re.DOTALL
            )

            if sub_heading_match:
                label = sub_heading_match.group(1).strip().rstrip(":")
                rest = sub_heading_match.group(2).strip()

                # Add as Heading 3
                doc.add_heading(label, level=3)

                # Add the body text
                if rest:
                    # Handle inline bold within the rest
                    add_rich_paragraph(doc, rest)
            else:
                # Regular paragraph — handle inline bold
                add_rich_paragraph(doc, para_text)


def add_rich_paragraph(doc: Document, text: str):
    """Add a paragraph with inline bold (**...**) formatting."""
    # Clean up any leading/trailing whitespace
    text = text.strip()
    if not text:
        return

    para = doc.add_paragraph()

    # Split by bold markers
    parts = re.split(r'\*\*(.*?)\*\*', text)

    for i, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        if i % 2 == 1:  # odd indices are bold content
            run.bold = True


def add_summary_section(doc: Document, summary_text: str):
    """Add the summary/assessment section with special formatting."""
    lines = summary_text.split("\n")
    heading_line = lines[0].lstrip("#").strip()

    # Use the summary heading style
    try:
        doc.add_paragraph(heading_line, style="SummaryHeading")
    except KeyError:
        doc.add_heading(heading_line, level=2)

    body = "\n".join(lines[1:]).strip()
    if body:
        add_body_paragraphs(doc, body)


# ─── Markdown Output ────────────────────────────────────────────────────────

def compile_markdown(jijnasa_episodes, jijnasa_summaries, janma_episodes, janma_summary):
    """Compile all content into a single Markdown string."""
    md = []
    md.append("# न्यायसुधा — Episode-wise Analytical Notes\n")
    md.append("**Pūrvapakṣa-Siddhānta Analysis of Nyāya Sudhā on Brahma Sūtra 1.1.1–1.1.2**\n")
    md.append("---\n")

    # Part I: Jijñāsādhikaraṇa
    md.append("## भाग I: जिज्ञासाधिकरणम् (BS 1.1.1)\n")

    for ep in jijnasa_episodes:
        md.append(f"### {ep['heading']}\n")
        md.append(f"{ep['body']}\n")
        md.append("---\n")

    for summary in jijnasa_summaries:
        if summary:
            md.append(f"\n{summary}\n")

    md.append("\n---\n")

    # Part II: Janmādhikaraṇa
    md.append("## भाग II: जन्माधिकरणम् (BS 1.1.2)\n")

    for ep in janma_episodes:
        md.append(f"### {ep['heading']}\n")
        md.append(f"{ep['body']}\n")
        md.append("---\n")

    if janma_summary:
        md.append(f"\n{janma_summary}\n")

    return "\n".join(md)


# ─── Main ────────────────────────────────────────────────────────────────────

def main():
    print("Extracting analysis from agent outputs...")

    # Collect all content
    jijnasa_episodes = []
    jijnasa_summaries = []
    janma_episodes = []
    janma_summary = ""

    for filename, part, order in SOURCE_FILES:
        filepath = TASK_DIR / filename
        if not filepath.exists():
            print(f"  WARNING: File not found: {filepath}")
            continue

        text = extract_analysis_text(filepath)
        if not text:
            print(f"  WARNING: No analysis found in {filename}")
            continue

        episodes, summary = parse_episodes(text)
        print(f"  {filename}: {len(episodes)} episodes, summary={'yes' if summary else 'no'}")

        if part == "jijnasa":
            jijnasa_episodes.extend(episodes)
            if summary:
                jijnasa_summaries.append(summary)
        elif part == "janma":
            janma_episodes.extend(episodes)
            if summary:
                janma_summary = summary

    total = len(jijnasa_episodes) + len(janma_episodes)
    print(f"\nTotal episodes collected: {total} "
          f"(Jijñāsā: {len(jijnasa_episodes)}, Janma: {len(janma_episodes)})")

    # ── Build DOCX ──

    print("\nBuilding DOCX...")
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.0)

    setup_styles(doc)

    # Title page
    doc.add_paragraph("न्यायसुधा — Episode-wise Analytical Notes", style="Title")
    doc.add_paragraph(
        "Pūrvapakṣa-Siddhānta Analysis of Nyāya Sudhā on Brahma Sūtra 1.1.1–1.1.2",
        style="Subtitle"
    )

    # Add author/project info
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_para.add_run(
        "Jijñāsā Adhikaraṇa Comparative Analysis Project\n"
        "Researcher: N. Kuvalaya Datta\n"
        "Supervisor: Prof. MLN Bhat (National Sanskrit University, Tirupati)\n"
        "Consultant: Dr. Vamśīkṛṣṇa Ghanapāṭhī"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    info_para.paragraph_format.space_after = Pt(30)

    # Horizontal rule
    doc.add_paragraph("─" * 60)

    # ── Part I: Jijñāsādhikaraṇa ──
    doc.add_heading("भाग I: जिज्ञासाधिकरणम् (BS 1.1.1)", level=1)
    subtitle_p = doc.add_paragraph()
    r = subtitle_p.add_run("Jijñāsādhikaraṇa — \"athāto brahma-jijñāsā\"")
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    for ep in jijnasa_episodes:
        doc.add_heading(ep["heading"], level=2)
        add_body_paragraphs(doc, ep["body"])

    # Jijñāsā summaries
    for summary in jijnasa_summaries:
        doc.add_page_break()
        add_summary_section(doc, summary)

    # ── Part II: Janmādhikaraṇa ──
    doc.add_page_break()
    doc.add_heading("भाग II: जन्माधिकरणम् (BS 1.1.2)", level=1)
    subtitle_p = doc.add_paragraph()
    r = subtitle_p.add_run("Janmādhikaraṇa — \"janmādy asya yataḥ\"")
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    for ep in janma_episodes:
        doc.add_heading(ep["heading"], level=2)
        add_body_paragraphs(doc, ep["body"])

    # Janma summary
    if janma_summary:
        doc.add_page_break()
        add_summary_section(doc, janma_summary)

    # Save DOCX
    docx_path = OUTPUT_DIR / "ns-analysis.docx"
    doc.save(str(docx_path))
    print(f"  Saved: {docx_path}")

    # ── Build Markdown ──

    print("Building Markdown...")
    md_content = compile_markdown(jijnasa_episodes, jijnasa_summaries,
                                  janma_episodes, janma_summary)
    md_path = OUTPUT_DIR / "ns-analysis.md"
    md_path.write_text(md_content, encoding="utf-8")
    print(f"  Saved: {md_path}")

    print(f"\nDone. {total} episodes compiled successfully.")


if __name__ == "__main__":
    main()

#!/usr/bin/env python3
"""Convert bs12-as.md to bs12-as-five-vyakhya.docx"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re


def main():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Default font for Normal style
    style = doc.styles['Normal']
    style.font.size = Pt(14)
    style.font.name = 'Noto Serif Devanagari'
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    # Heading styles
    for level, size in [(1, 18), (2, 16), (3, 15)]:
        h = doc.styles[f'Heading {level}']
        h.font.size = Pt(size)
        h.font.name = 'Noto Serif Devanagari'
        h.font.bold = True
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)

    with open('bs12-as.md', 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Skip empty lines
        if not line:
            i += 1
            continue

        # Skip horizontal rules
        if line == '---':
            i += 1
            continue

        # Skip Source: lines
        if line.startswith('Source:'):
            i += 1
            continue

        # Heading 1 (top-level: # )
        if line.startswith('# ') and not line.startswith('## '):
            doc.add_heading(line[2:], level=1)
            i += 1
            continue

        # Heading 2 (## )
        if line.startswith('## '):
            doc.add_heading(line[3:], level=2)
            i += 1
            continue

        # Heading 3 (### )
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
            i += 1
            continue

        # Bold markers (**text**) in prateeka lines
        if line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line.strip('*'))
            run.bold = True
            i += 1
            continue

        # Regular paragraph
        # Clean up markdown bold markers inline
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
        doc.add_paragraph(text)
        i += 1

    outfile = 'bs12-as-five-vyakhya.docx'
    doc.save(outfile)
    print(f'Saved {outfile}')


if __name__ == '__main__':
    main()

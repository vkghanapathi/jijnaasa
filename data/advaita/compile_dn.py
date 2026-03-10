#!/usr/bin/env python3
"""
compile_dn.py
-------------
Combines the two DN episode files into a single Advaita reference document:
  1. 04-dn-episodes.docx
  2. 04-dn-episodes.md
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = Path("/Users/dharmaposhanam/Documents/GitHub/Jijnaasa/data/advaita")

SOURCE_FILES = [
    BASE_DIR / "dn-episodes-jijnasa.md",
    BASE_DIR / "dn-episodes-janmadya.md",
]

OUT_DOCX = BASE_DIR / "04-dn-episodes.docx"
OUT_MD = BASE_DIR / "04-dn-episodes.md"


def setup_styles(doc):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Cambria"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    title = styles["Title"]
    title.font.name = "Cambria"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x1A, 0x1A, 0x5C)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    h1 = styles["Heading 1"]
    h1.font.name = "Cambria"
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

    h2 = styles["Heading 2"]
    h2.font.name = "Cambria"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x1A, 0x1A, 0x5C)

    h3 = styles["Heading 3"]
    h3.font.name = "Cambria"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_rich_paragraph(doc, text):
    text = text.strip()
    if not text:
        return
    para = doc.add_paragraph()
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        if i % 2 == 1:
            run.bold = True


def md_to_docx(doc, lines):
    """Convert markdown lines to docx content."""
    for line in lines:
        s = line.rstrip()

        if not s:
            continue

        # Skip HR
        if re.match(r'^---+$', s):
            continue

        # Title (#)
        m = re.match(r'^#\s+(.+)$', s)
        if m and not re.match(r'^##', s):
            doc.add_paragraph(m.group(1), style="Title")
            continue

        # H1 (##)
        m = re.match(r'^##\s+(?!#)(.+)$', s)
        if m:
            doc.add_heading(m.group(1), level=1)
            continue

        # H2 (###)
        m = re.match(r'^###\s+(?!#)(.+)$', s)
        if m:
            doc.add_heading(m.group(1), level=2)
            continue

        # H3 (####)
        m = re.match(r'^####\s+(?!#)(.+)$', s)
        if m:
            doc.add_heading(m.group(1), level=3)
            continue

        # Bold section markers
        m = re.match(r'^\*\*(.+?)\*\*\s*$', s)
        if m:
            doc.add_heading(m.group(1), level=3)
            continue

        # Table lines — skip formatting headers
        if s.startswith('|') and ':---' in s:
            continue
        if s.startswith('|'):
            text = s.strip('| ').replace('|', ' — ')
            add_rich_paragraph(doc, text)
            continue

        # Regular text
        add_rich_paragraph(doc, s)


def main():
    print("Compiling DN episodes...")

    # Combine markdown
    combined_md = []
    combined_md.append("# ध्वान्तनिरासः — प्रकरणानुक्रमणिका")
    combined_md.append("")
    combined_md.append("**न्यायसुधोक्तदूषणानाम् उद्धारः (Advaita Counter-refutation of Nyāya Sudhā)**")
    combined_md.append("")
    combined_md.append("---")
    combined_md.append("")

    for filepath in SOURCE_FILES:
        if not filepath.exists():
            print(f"  WARNING: {filepath.name} not found")
            continue
        text = filepath.read_text(encoding='utf-8')
        lines = text.splitlines()
        combined_md.extend(lines)
        combined_md.append("")
        print(f"  {filepath.name}: {len(lines)} lines")

    # Write combined markdown
    OUT_MD.write_text('\n'.join(combined_md), encoding='utf-8')
    print(f"  Saved: {OUT_MD}")

    # Build DOCX
    print("Building DOCX...")
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.0)

    setup_styles(doc)

    md_to_docx(doc, combined_md)

    doc.save(str(OUT_DOCX))
    print(f"  Saved: {OUT_DOCX}")
    print("Done.")


if __name__ == "__main__":
    main()

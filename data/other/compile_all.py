#!/usr/bin/env python3
"""
compile_all.py
--------------
Produces two docx files in data/other/:
  1. 03-ns-combined.docx — Original NS Sanskrit text + English analysis per episode
  2. 05-ns-review.docx   — Copy of Sanskrit analysis (already compiled)

Reads from:
  - data/dvaita/bs-ns12.md     (original NS text, ~8983 lines)
  - data/dvaita/ns-analysis.md (English analysis, ~1098 lines)
  - data/dvaita/05-ns-review.docx (Sanskrit, copy)
"""

import re
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ─── Paths ──────────────────────────────────────────────────────────────────

DVAITA_DIR = Path("/Users/dharmaposhanam/Documents/GitHub/Jijnaasa/data/dvaita")
OTHER_DIR = Path("/Users/dharmaposhanam/Documents/GitHub/Jijnaasa/data/other")

NS_SOURCE = DVAITA_DIR / "bs-ns12.md"
NS_ANALYSIS = DVAITA_DIR / "ns-analysis.md"
SANSKRIT_DOCX = DVAITA_DIR / "05-ns-review.docx"

OUT_COMBINED = OTHER_DIR / "03-ns-combined.docx"
OUT_SANSKRIT = OTHER_DIR / "05-ns-review.docx"

# ─── Parse original NS text ────────────────────────────────────────────────

# Episode heading in bs-ns12.md: # N. title (M sentences)
NS_EPISODE_RE = re.compile(r'^#\s+(\d+)\.\s+(.+?)(?:\s+\(\d+\s+sentences\))?\s*$')
NS_SECTION_RE = re.compile(r'^##\s+(.+)$')
NS_SUBHEAD_RE = re.compile(r'^##\s+(.+)$')


def parse_ns_source(filepath):
    """Parse bs-ns12.md into episodes with their source text."""
    lines = filepath.read_text(encoding='utf-8').splitlines()

    episodes = {}  # num -> {'title': ..., 'part': ..., 'lines': [...]}
    current_ep = None
    current_part = None
    in_toc = False

    for line in lines:
        s = line.rstrip()

        # Skip TOC section
        if s.startswith('# विषयानुक्रमणिका'):
            in_toc = True
            continue
        if in_toc:
            # TOC ends at first section heading or episode heading
            if NS_EPISODE_RE.match(s):
                in_toc = False
            elif s.startswith('## भाग'):
                current_part = s.lstrip('# ').strip()
                continue
            else:
                continue

        # Section heading (भाग)
        m = re.match(r'^##?\s+भाग\s+(.+)$', s)
        if m:
            current_part = s.lstrip('# ').strip()
            continue

        # Episode heading
        m = NS_EPISODE_RE.match(s)
        if m:
            if current_ep:
                # Trim trailing blanks
                while current_ep['lines'] and not current_ep['lines'][-1].strip():
                    current_ep['lines'].pop()

            num = int(m.group(1))
            title = m.group(2).strip()
            # Remove sentence count from title if present
            title = re.sub(r'\s*\(\d+\s+sentences?\)\s*$', '', title)
            current_ep = {
                'title': title,
                'part': current_part,
                'lines': [],
            }
            episodes[num] = current_ep
            continue

        # Accumulate text
        if current_ep is not None:
            # Skip sub-headings like ## न्यायसुधा
            if s.startswith('## '):
                continue
            # Skip table formatting lines
            if s.startswith('|') and ':---' in s:
                continue
            current_ep['lines'].append(s)

    # Trim last episode
    if current_ep:
        while current_ep['lines'] and not current_ep['lines'][-1].strip():
            current_ep['lines'].pop()

    return episodes


# ─── Parse English analysis ────────────────────────────────────────────────

# Episode heading in ns-analysis.md: ### Episode N: Title
ANALYSIS_EP_RE = re.compile(r'^###\s+Episode\s+(\d+)\s*[:\.\—–-]\s*(.+)$')
ANALYSIS_PART_RE = re.compile(r'^##\s+(.+)$')


def parse_analysis(filepath):
    """Parse ns-analysis.md into episodes."""
    lines = filepath.read_text(encoding='utf-8').splitlines()

    episodes = {}
    current_ep = None
    current_part = None

    for line in lines:
        s = line.rstrip()

        # Part heading
        m = ANALYSIS_PART_RE.match(s)
        if m and not s.startswith('###'):
            current_part = m.group(1).strip()
            continue

        # Episode heading
        m = ANALYSIS_EP_RE.match(s)
        if m:
            if current_ep:
                while current_ep['lines'] and not current_ep['lines'][-1].strip():
                    current_ep['lines'].pop()

            num = int(m.group(1))
            title = m.group(2).strip()
            current_ep = {
                'title': title,
                'part': current_part,
                'lines': [],
            }
            episodes[num] = current_ep
            continue

        if current_ep is not None:
            # Skip horizontal rules
            if re.match(r'^---+$', s):
                continue
            current_ep['lines'].append(s)

    if current_ep:
        while current_ep['lines'] and not current_ep['lines'][-1].strip():
            current_ep['lines'].pop()

    return episodes


# ─── DOCX Building ─────────────────────────────────────────────────────────

def setup_styles(doc):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Cambria"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    title = styles["Title"]
    title.font.name = "Cambria"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x1A, 0x1A, 0x5C)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    h1 = styles["Heading 1"]
    h1.font.name = "Cambria"
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

    h2 = styles["Heading 2"]
    h2.font.name = "Cambria"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x1A, 0x1A, 0x5C)

    h3 = styles["Heading 3"]
    h3.font.name = "Cambria"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_rich_paragraph(doc, text):
    text = text.strip()
    if not text:
        return
    para = doc.add_paragraph()
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        if i % 2 == 1:
            run.bold = True


def add_source_text(doc, lines):
    """Add original Sanskrit source text with distinct formatting."""
    # Add a visual separator
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sep.add_run("── मूलपाठः ──")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x44, 0x00)
    run.bold = True

    for line in lines:
        s = line.strip()
        if not s:
            continue
        # Table lines
        if s.startswith('|'):
            text = s.strip('| ').replace('|', ' — ')
            if ':---' in text:
                continue
            para = doc.add_paragraph(text)
            para.style.font.size = Pt(10)
            continue
        # Regular Sanskrit text
        para = doc.add_paragraph(s)
        for run in para.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # End separator
    sep2 = doc.add_paragraph()
    sep2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sep2.add_run("── ── ──")
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x88, 0x44, 0x00)


def add_analysis_text(doc, lines):
    """Add English analysis text."""
    doc.add_heading("विश्लेषणम् (Analysis)", level=3)

    for line in lines:
        s = line.strip()
        if not s:
            continue

        # Bold section markers
        m = re.match(r'^\*\*(.+?)\*\*\s*[:\s]*(.*)', s, re.DOTALL)
        if m:
            label = m.group(1).strip().rstrip(':')
            rest = m.group(2).strip()
            doc.add_heading(label, level=3)
            if rest:
                add_rich_paragraph(doc, rest)
            continue

        add_rich_paragraph(doc, s)


def build_combined_docx(ns_episodes, analysis_episodes):
    """Build the combined English docx."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.0)

    setup_styles(doc)

    # Title
    doc.add_paragraph("न्यायसुधा — समग्रविश्लेषणम्", style="Title")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Original Text + Episode-wise Analysis\n"
        "Brahma Sūtra 1.1.1–1.1.2"
    )
    r.font.size = Pt(13)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # Info
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_para.add_run(
        "Jijñāsā Adhikaraṇa Comparative Analysis Project\n"
        "Researcher: N. Kuvalaya Datta\n"
        "Supervisor: Prof. MLN Bhat (NSU, Tirupati)\n"
        "Consultant: Dr. Vamśīkṛṣṇa Ghanapāṭhī"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph("─" * 60)

    # Determine all episode numbers
    all_nums = sorted(set(list(ns_episodes.keys()) + list(analysis_episodes.keys())))

    # Detect part breaks
    current_part = None

    for num in all_nums:
        ns_ep = ns_episodes.get(num)
        an_ep = analysis_episodes.get(num)

        # Part heading
        part = None
        if ns_ep:
            part = ns_ep.get('part')
        if not part and an_ep:
            part = an_ep.get('part')

        if part and part != current_part:
            current_part = part
            doc.add_page_break()
            doc.add_heading(part, level=1)

        # Episode heading
        title = ns_ep['title'] if ns_ep else (an_ep['title'] if an_ep else f"Episode {num}")
        doc.add_heading(f"{num}. {title}", level=2)

        # Original source text
        if ns_ep and ns_ep['lines']:
            add_source_text(doc, ns_ep['lines'])

        # Analysis
        if an_ep and an_ep['lines']:
            add_analysis_text(doc, an_ep['lines'])

    doc.save(str(OUT_COMBINED))
    print(f"  Saved: {OUT_COMBINED}")
    return len(all_nums)


# ─── Main ──────────────────────────────────────────────────────────────────

def main():
    # Task 1: Combined English
    print("Parsing original NS text...")
    ns_episodes = parse_ns_source(NS_SOURCE)
    print(f"  Found {len(ns_episodes)} NS episodes")

    print("Parsing English analysis...")
    analysis_episodes = parse_analysis(NS_ANALYSIS)
    print(f"  Found {len(analysis_episodes)} analysis episodes")

    print("\nBuilding combined DOCX...")
    count = build_combined_docx(ns_episodes, analysis_episodes)
    print(f"  {count} episodes in combined document")

    # Task 2: Copy Sanskrit
    print("\nCopying Sanskrit analysis...")
    shutil.copy2(str(SANSKRIT_DOCX), str(OUT_SANSKRIT))
    print(f"  Saved: {OUT_SANSKRIT}")

    print("\nDone.")


if __name__ == "__main__":
    main()

#!/usr/bin/env python3
"""
compile_telugu.py
-----------------
Combines 3 Telugu translation files into:
  1. 03-ns-tl.docx (Word document)
  2. 03-ns-tl.md   (Markdown reference)
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OTHER_DIR = Path("/Users/dharmaposhanam/Documents/GitHub/Jijnaasa/data/other")

SOURCE_FILES = [
    OTHER_DIR / "telugu_eps_01_30.md",
    OTHER_DIR / "telugu_eps_31_60.md",
    OTHER_DIR / "telugu_eps_61_end.md",
]

OUT_DOCX = OTHER_DIR / "03-ns-tl.docx"
OUT_MD = OTHER_DIR / "03-ns-tl.md"


def setup_styles(doc):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Cambria"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    title = styles["Title"]
    title.font.name = "Cambria"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x1A, 0x1A, 0x5C)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    h1 = styles["Heading 1"]
    h1.font.name = "Cambria"
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

    h2 = styles["Heading 2"]
    h2.font.name = "Cambria"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x1A, 0x1A, 0x5C)

    h3 = styles["Heading 3"]
    h3.font.name = "Cambria"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_rich_paragraph(doc, text):
    text = text.strip()
    if not text:
        return
    para = doc.add_paragraph()
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        if i % 2 == 1:
            run.bold = True


def md_to_docx(doc, lines):
    """Convert markdown lines to docx content."""
    for line in lines:
        s = line.rstrip()

        if not s:
            continue

        # Skip HR
        if re.match(r'^---+$', s):
            continue

        # H1 (#) - but not ## or ###
        m = re.match(r'^#\s+(?!#)(.+)$', s)
        if m:
            doc.add_paragraph(m.group(1), style="Title")
            continue

        # H2 (##) - but not ###
        m = re.match(r'^##\s+(?!#)(.+)$', s)
        if m:
            doc.add_heading(m.group(1), level=1)
            continue

        # H3 (###) - but not ####
        m = re.match(r'^###\s+(?!#)(.+)$', s)
        if m:
            doc.add_heading(m.group(1), level=2)
            continue

        # H4 (####)
        m = re.match(r'^####\s+(.+)$', s)
        if m:
            doc.add_heading(m.group(1), level=3)
            continue

        # Bold section markers like **పూర్వపక్షము:**
        m = re.match(r'^\*\*(.+?)\*\*\s*$', s)
        if m:
            doc.add_heading(m.group(1), level=3)
            continue

        # Bold label with content: **Label:** content
        m = re.match(r'^\*\*(.+?)\*\*\s*[:\—–-]\s*(.+)', s)
        if m:
            doc.add_heading(m.group(1).rstrip(':'), level=3)
            add_rich_paragraph(doc, m.group(2))
            continue

        # Regular text
        add_rich_paragraph(doc, s)


def main():
    print("Compiling Telugu translation...")

    # Combine markdown
    combined = []
    combined.append("# న్యాయసుధా — ప్రకరణవారీ విశ్లేషణ టిప్పణులు")
    combined.append("")
    combined.append("**బ్రహ్మసూత్ర 1.1.1–1.1.2 పై న్యాయసుధా యొక్క పూర్వపక్ష-సిద్ధాంత విశ్లేషణము**")
    combined.append("")
    combined.append("---")
    combined.append("")

    total_lines = 0
    for filepath in SOURCE_FILES:
        if not filepath.exists():
            print(f"  WARNING: {filepath.name} not found")
            continue
        text = filepath.read_text(encoding='utf-8')
        lines = text.splitlines()

        # Skip the title line of subsequent files to avoid duplication
        start = 0
        for i, line in enumerate(lines):
            if line.startswith('# '):
                start = i + 1
                break

        file_lines = lines[start:]
        combined.extend(file_lines)
        combined.append("")
        total_lines += len(file_lines)
        print(f"  {filepath.name}: {len(file_lines)} lines")

    # Write combined markdown
    OUT_MD.write_text('\n'.join(combined), encoding='utf-8')
    print(f"  Saved: {OUT_MD}")

    # Build DOCX
    print("Building DOCX...")
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.0)

    setup_styles(doc)

    # Add project info
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_para.add_run(
        "జిజ్ఞాసా అధికరణ తులనాత్మక విశ్లేషణ ప్రకల్పము\n"
        "అన్వేషకుడు: ఎన్. కువలయ దత్తః\n"
        "మార్గదర్శకుడు: ప్రొ. ఎం.ఎల్.ఎన్. భట్ (రాష్ట్రీయ సంస్కృత విశ్వవిద్యాలయము, తిరుపతి)\n"
        "పరామర్శదాత: డా. వంశీకృష్ణ ఘనపాఠీ"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    info_para.paragraph_format.space_after = Pt(20)

    md_to_docx(doc, combined)

    doc.save(str(OUT_DOCX))
    print(f"  Saved: {OUT_DOCX}")

    print(f"\nDone. {total_lines} total lines compiled.")


if __name__ == "__main__":
    main()

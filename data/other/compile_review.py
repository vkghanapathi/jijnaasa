#!/usr/bin/env python3
"""
compile_review.py
-----------------
Master compilation script for the Jijñāsā Adhikaraṇa comprehensive review.
Produces a ~320-page DOCX combining all source layers:

  Part A: Advaita Foundational Texts (Bhāṣya + 5 Vyākhyās + Ānandagiri)
  Part B: Jijñāsādhikaraṇam (BS 1.1.1) — 83 episodes, layered
  Part C: Janmādhikaraṇam (BS 1.1.2) — 6 episodes, layered
  Appendices: DN cross-reference, unmatched DN prakāraṇas, bibliography

Each episode in Parts B & C contains:
  - NS Original Sanskrit (full, from bs-ns12.md)
  - Anuvyākhyāna verse
  - Sanskrit Analytical Review (from 05-ns-review.md)
  - DN Counter-refutation (where mapped)
  - Samīkṣā slot (blank reviewer notes)

Reads from:
  data/dvaita/bs-ns12.md, ns-analysis.md, 05-ns-review.md
  data/advaita/dn-episodes-jijnasa.md, dn-episodes-janmadya.md
  data/advaita/bs12-as.md, 01-bs12-anandagiri-sutramuktavali.md

Output:
  data/other/05-ns-review.docx
  data/other/05-ns-review.md
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ═══════════════════════════════════════════════════════════════════════════
# PATHS
# ═══════════════════════════════════════════════════════════════════════════

BASE = Path("/Users/dharmaposhanam/Documents/GitHub/Jijnaasa")
DVAITA = BASE / "data/dvaita"
ADVAITA = BASE / "data/advaita"
OTHER = BASE / "data/other"

NS_SOURCE_FILE = DVAITA / "bs-ns12.md"
NS_REVIEW_FILE = DVAITA / "05-ns-review.md"
DN_JIJNASA_FILE = ADVAITA / "dn-episodes-jijnasa.md"
DN_JANMADYA_FILE = ADVAITA / "dn-episodes-janmadya.md"
BHASYA_FILE = ADVAITA / "bs12-as.md"
SUTRAMUKTAVALI_FILE = ADVAITA / "01-bs12-anandagiri-sutramuktavali.md"

OUT_DOCX = OTHER / "05-ns-review.docx"
OUT_MD = OTHER / "05-ns-review.md"

# ═══════════════════════════════════════════════════════════════════════════
# COLORS
# ═══════════════════════════════════════════════════════════════════════════

DARK_BLUE = RGBColor(0x1A, 0x1A, 0x5C)
DARK_RED = RGBColor(0x8B, 0x00, 0x00)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MED_GRAY = RGBColor(0x66, 0x66, 0x66)
BROWN = RGBColor(0x88, 0x44, 0x00)
DARK_GREEN = RGBColor(0x00, 0x66, 0x00)
SAFFRON = RGBColor(0xCC, 0x66, 0x00)
LIGHT_GRAY = RGBColor(0x99, 0x99, 0x99)
DEEP_PURPLE = RGBColor(0x4B, 0x00, 0x82)  # Vyākhyā citations

# ═══════════════════════════════════════════════════════════════════════════
# DN → NS EPISODE MAPPING
# ═══════════════════════════════════════════════════════════════════════════
# Maps DN prakāraṇa number → list of NS episode numbers it addresses.
# Built from scholarly topic analysis.

DN_JIJNASA_MAP = {
    1: [22],                    # Śārīrakaśabdārthavicāra
    2: [26, 27, 28],            # Athaśabdārtha — maṅgalasamanvaya
    3: [28],                    # Vākyārthacandrikā anyonyāśrayadoṣa
    4: [39, 40, 41, 43, 46],    # Oṃkāre jīvabrahmaikya
    5: [58, 59],                # Anādibhāvarūpājñāna nivṛtti (Vivaraṇa)
    6: [59],                    # Anādibhāvasya vināśasambhava
    7: [54, 57, 61],            # Yauktikajñānasya ajñānanivartakatva
    8: [76, 77, 78, 79, 80],    # Anirvacanīyatvavicāra
    9: [30, 33],                # Bhāmatīdūṣaṇanirāsa — anuvādatva
    10: [53, 54, 55, 56, 57, 62, 63, 64, 65, 66],  # Mukti-bandha-mithyātva
    11: [64],                   # Satya-viṣa-nivṛtti-dṛṣṭānta
    12: [68, 71],               # Neha nānā — mithyātvasādhakatva
    13: [67],                   # Itaretarādhyāsa — mithyātvasādhakatva
    14: [35, 36],               # Ahamarthavicāra — ahaṅkārātmabheda
    15: [35, 36],               # Suṣupti — ahamanubhava
    16: [69, 70, 72, 73, 74, 75],  # Mithyābhūtaśruti sādhakatva
    17: [68, 71],               # Neha nānā — bhedaniṣedha
    18: [73, 74],               # Sākṣiṇi ajñānasiddhi — itaretarāśraya
    19: [69, 72, 73],           # Māyāvādinaḥ kathakatā — asataḥ sādhakatva
    20: [76, 77, 78, 79, 80],   # Sad-asad-vilakṣaṇa daśavidhadūṣaṇa
    21: [78, 79],               # Jñānabādhyatva — mithyātvahetu
    22: [80],                   # Anirvacanīyatvasya svarūpa
    23: [76, 77],               # Sattvatraividhya — anumāna
    24: [81, 82, 83],           # Śuktirūpyata — anirvacanīyakhyāti
    25: [69, 70],               # Asaccen na pratīyeta — pratītatvahetu
    26: [76, 77],               # Sattvatraividhya — anavasthā
    27: [58, 59, 60],           # Brahmaṇi āvaraṇa — ajñānāśraya
    28: [22, 35, 36, 67],       # Tattvamasi — pratītārthatva
}

DN_JANMADYA_MAP = {
    1: [1, 2, 3],               # Upādānatvatraividhya
    2: [3],                     # Dhyānaṃ jñānaṃ ca
    3: [3, 6],                  # Advaitaśrutyabhiprāyatva
    4: [6],                     # Vivartavāda
    5: [2, 3],                  # Taṭasthalakṣaṇatve lakṣaṇāśrayaṇa
    6: [3, 4],                  # Tattvamasi — jīvabrahmaikya
    7: [3, 4],                  # Virodhyākāraparityāga
    8: [1, 3, 4],               # Nirguṇatvavicāra
    9: [4],                     # Satyaṃ jñānam — ākhaṇḍārtha
    10: [4],                    # Satyādipadānāṃ vyāvṛtti
    11: [4, 5],                 # Nirviśeṣaṃ brahma
    12: [4, 5],                 # Viśeṣakalpanā — Madhvavaiśeṣika
    13: [6],                    # Brahmaṇo jagatakāraṇatva
}

# Reverse map: for each NS episode, which DN prakāraṇas address it
def _build_reverse_map(dn_map):
    rev = {}
    for dn_num, ns_nums in dn_map.items():
        for ns_num in ns_nums:
            rev.setdefault(ns_num, []).append(dn_num)
    return rev

# ═══════════════════════════════════════════════════════════════════════════
# REGEX PATTERNS
# ═══════════════════════════════════════════════════════════════════════════

NS_EPISODE_RE = re.compile(
    r'^#\s+(\d+)\.\s+(.+?)(?:\s+\(\d+\s+sentences?\))?\s*$'
)
REVIEW_EP_RE = re.compile(
    r'^###\s+प्रकरणम्\s+(\d+)\s*[:\.\—–-]\s*(.+)$'
)
DN_EP_RE = re.compile(
    r'^###\s+प्रकरणम्\s+(\d+)\s*[:\.\—–-]\s*(.+)$'
)
SECTION_BOLD_RE = re.compile(r'^\*\*(.+?)\*\*\s*$')
PART_RE = re.compile(r'^##\s+(.+)$')
HR_RE = re.compile(r'^---+$')
BHASYA_SECTION_RE = re.compile(r'^##\s+(.+)$')
VYAKHYA_RE = re.compile(r'^###\s+(.+)$')

# ═══════════════════════════════════════════════════════════════════════════
# MODULE 1: PARSERS
# ═══════════════════════════════════════════════════════════════════════════


def parse_ns_source(filepath):
    """Parse bs-ns12.md into episodes dict keyed by (part, num).
    Returns: {'jijnasa': {1: {...}, 2: {...}, ...}, 'janma': {1: {...}, ...}}
    """
    lines = filepath.read_text(encoding='utf-8').splitlines()
    episodes = {'jijnasa': {}, 'janma': {}}
    current_ep = None
    current_part = 'jijnasa'
    in_toc = False

    for line in lines:
        s = line.rstrip()

        # Skip TOC — do NOT track part labels inside TOC
        if s.startswith('# विषयानुक्रमणिका'):
            in_toc = True
            continue
        if in_toc:
            if NS_EPISODE_RE.match(s):
                in_toc = False
                current_part = 'jijnasa'  # Body starts with Part 1
            else:
                continue

        # Part boundary (# भाग or ## भाग)
        m = re.match(r'^#{1,2}\s+भाग\s+(.+)$', s)
        if m:
            txt = m.group(1)
            if 'जन्मा' in txt or '१.१.२' in txt or '1.1.2' in txt or '२' in txt:
                current_part = 'janma'
            else:
                current_part = 'jijnasa'
            continue

        # Episode heading
        m = NS_EPISODE_RE.match(s)
        if m:
            if current_ep:
                while current_ep['lines'] and not current_ep['lines'][-1].strip():
                    current_ep['lines'].pop()
            num = int(m.group(1))
            title = re.sub(r'\s*\(\d+\s+sentences?\)\s*$', '', m.group(2).strip())
            current_ep = {'title': title, 'part': current_part, 'lines': []}
            episodes[current_part][num] = current_ep
            continue

        # Accumulate text
        if current_ep is not None:
            if s.startswith('## '):
                continue
            if s.startswith('|') and ':---' in s:
                continue
            current_ep['lines'].append(s)

    if current_ep:
        while current_ep['lines'] and not current_ep['lines'][-1].strip():
            current_ep['lines'].pop()

    return episodes


def parse_review(filepath):
    """Parse 05-ns-review.md into structured episodes.
    Returns: {'jijnasa': {1: {...}, ...}, 'janma': {1: {...}, ...}}
    """
    lines = filepath.read_text(encoding='utf-8').splitlines()
    episodes = {'jijnasa': {}, 'janma': {}}
    current_part = 'jijnasa'
    current_ep = None
    current_section_name = None
    current_section_lines = []

    def flush_section():
        nonlocal current_section_name, current_section_lines
        if current_ep and current_section_name:
            current_ep['sections'].append({
                'name': current_section_name,
                'text': '\n'.join(current_section_lines).strip(),
            })
        current_section_name = None
        current_section_lines = []

    def flush_episode():
        flush_section()

    for line in lines:
        s = line.rstrip()

        # Part heading
        m = PART_RE.match(s)
        if m and not s.startswith('###'):
            txt = m.group(1)
            if 'जन्मा' in txt or '१.१.२' in txt or '1.1.2' in txt:
                current_part = 'janma'
            continue

        # Episode heading
        m = REVIEW_EP_RE.match(s)
        if m:
            flush_episode()
            num = int(m.group(1))
            title = m.group(2).strip()
            current_ep = {
                'num': num,
                'title': title,
                'part': current_part,
                'sections': [],
                'intro': [],
            }
            episodes[current_part][num] = current_ep
            continue

        if current_ep is None:
            continue

        # HR — skip
        if HR_RE.match(s):
            continue

        # Section bold heading
        m = SECTION_BOLD_RE.match(s)
        if m:
            flush_section()
            current_section_name = m.group(1).strip()
            current_section_lines = []
            continue

        # Accumulate
        if current_section_name:
            current_section_lines.append(s)
        elif s.strip():
            current_ep['intro'].append(s)

    flush_episode()
    return episodes


def parse_dn_episodes(filepath):
    """Parse a DN episode file into dict keyed by prakāraṇa number.
    Each entry has: title, vishaya, vyakhya, purvapaksha, khandana, siddhanta, raw.
    """
    if not filepath.exists():
        return {}
    lines = filepath.read_text(encoding='utf-8').splitlines()
    episodes = {}
    current_ep = None
    current_section = None
    current_lines = []

    def flush_section():
        nonlocal current_section, current_lines
        if current_ep and current_section:
            current_ep[current_section] = '\n'.join(current_lines).strip()
        current_section = None
        current_lines = []

    SECTION_MAP = {
        'विषयः': 'vishaya',
        'व्याख्या': 'vyakhya',
        'पूर्वपक्षः': 'purvapaksha',
        'खण्डनम्': 'khandana',
        'सिद्धान्तः': 'siddhanta',
    }

    for line in lines:
        s = line.rstrip()

        # Episode heading
        m = DN_EP_RE.match(s)
        if m:
            flush_section()
            num = int(m.group(1))
            title = m.group(2).strip()
            current_ep = {
                'num': num,
                'title': title,
                'vishaya': '',
                'vyakhya': '',
                'purvapaksha': '',
                'khandana': '',
                'siddhanta': '',
                'raw': [],
            }
            episodes[num] = current_ep
            continue

        if current_ep is None:
            continue

        # HR — skip
        if HR_RE.match(s):
            continue

        current_ep['raw'].append(s)

        # Bold section
        m = re.match(r'^\*\*(.+?)\*\*', s)
        if m:
            label = m.group(1).strip()
            # Strip trailing line ref like (line 482-495)
            label_clean = re.sub(r'\s*\(line.*?\)\s*$', '', label).strip()
            key = SECTION_MAP.get(label_clean)
            if key:
                flush_section()
                current_section = key
                # Capture text after the bold marker on same line
                rest = re.sub(r'^\*\*.+?\*\*\s*', '', s)
                rest = re.sub(r'\(line\s+\d+[–-]\d+\)\s*:\s*', '', rest)
                if rest.strip():
                    current_lines.append(rest.strip())
                continue

        if current_section:
            current_lines.append(s)

    flush_section()
    return episodes


def parse_bhasya_sections(filepath):
    """Parse bs12-as.md into major sections with vyākhyā sub-sections.
    Returns list of dicts: [{'heading': ..., 'content': [...], 'vyakhyas': [...]}]
    """
    if not filepath.exists():
        return []
    lines = filepath.read_text(encoding='utf-8').splitlines()
    sections = []
    current_section = None
    current_vyakhya = None
    current_lines = []

    def flush_vyakhya():
        nonlocal current_vyakhya, current_lines
        if current_section and current_vyakhya:
            current_section['vyakhyas'].append({
                'name': current_vyakhya,
                'lines': current_lines,
            })
        elif current_section and current_lines:
            current_section['content'].extend(current_lines)
        current_vyakhya = None
        current_lines = []

    for line in lines:
        s = line.rstrip()

        # Major section (##)
        m = re.match(r'^##\s+(?!#)(.+)$', s)
        if m:
            flush_vyakhya()
            if current_section:
                sections.append(current_section)
            current_section = {
                'heading': m.group(1).strip(),
                'content': [],
                'vyakhyas': [],
            }
            continue

        # Vyākhyā subsection (###)
        m = re.match(r'^###\s+(?!#)(.+)$', s)
        if m:
            flush_vyakhya()
            current_vyakhya = m.group(1).strip()
            current_lines = []
            continue

        current_lines.append(s)

    flush_vyakhya()
    if current_section:
        sections.append(current_section)

    return sections


def read_raw_text(filepath):
    """Read a file as raw text."""
    if not filepath.exists():
        return ""
    return filepath.read_text(encoding='utf-8')


def parse_vyakhya_subsections(bhasya_sections):
    """Parse vyākhyā into granular sub-sections for per-episode matching.

    Returns list of dicts:
        {'vyakhya': name, 'sutra': 'adhyasa'|'sutra1'|'sutra2',
         'heading': bold header, 'text': content}
    """
    subsections = []
    for sec in bhasya_sections:
        heading = sec['heading']
        if 'सूत्रम् १' in heading:
            sutra = 'sutra1'
        elif 'सूत्रम् २' in heading:
            sutra = 'sutra2'
        elif 'अध्यास' in heading:
            sutra = 'adhyasa'
        else:
            sutra = 'adhyasa'

        for vy in sec['vyakhyas']:
            current_heading = None
            current_lines = []

            for line in vy['lines']:
                s = line.strip()
                m = SECTION_BOLD_RE.match(s)
                if m:
                    if current_heading and current_lines:
                        text = '\n'.join(current_lines).strip()
                        if text:
                            subsections.append({
                                'vyakhya': vy['name'],
                                'sutra': sutra,
                                'heading': current_heading,
                                'text': text,
                            })
                    current_heading = m.group(1).strip()
                    current_lines = []
                else:
                    current_lines.append(s)

            if current_heading and current_lines:
                text = '\n'.join(current_lines).strip()
                if text:
                    subsections.append({
                        'vyakhya': vy['name'],
                        'sutra': sutra,
                        'heading': current_heading,
                        'text': text,
                    })

            # If no bold sub-sections, treat the whole vyākhyā as one block
            if not any(ss['vyakhya'] == vy['name'] and ss['sutra'] == sutra
                       for ss in subsections):
                full = '\n'.join(vy['lines']).strip()
                if full:
                    subsections.append({
                        'vyakhya': vy['name'],
                        'sutra': sutra,
                        'heading': vy['name'],
                        'text': full,
                    })

    return subsections


# Sanskrit keyword stems for topic matching (4+ chars for compound splitting)
_STOP_STEMS = {
    'इति', 'तथा', 'यथा', 'अत्र', 'अतः', 'तस्य', 'तस्मात्', 'एवम्',
    'यस्य', 'तेन', 'अयम्', 'इदम्', 'कथम्', 'किम्', 'सर्व',
}


def _extract_stems(text, min_len=4):
    """Extract significant Sanskrit stems from text for matching."""
    # Split on spaces, punctuation, virama combinations
    tokens = re.split(r'[\s।॥,\-–—:()]+', text)
    stems = set()
    for t in tokens:
        t = t.strip('।॥*')
        if len(t) >= min_len and t not in _STOP_STEMS:
            stems.add(t)
            # Also add shorter stem (first 5 chars) for compound matching
            if len(t) >= 6:
                stems.add(t[:5])
    return stems


def find_vyakhya_for_episode(ep_title, ep_lines, subsections, sutra_filter,
                              max_per_vyakhya=2):
    """Find relevant vyākhyā sub-sections for an NS episode.

    Returns dict: {vyakhya_name: [{'heading': ..., 'text': ...}, ...]}
    """
    # Extract keywords from episode title + first 30 content lines
    title_stems = _extract_stems(ep_title, min_len=4)
    content_sample = ' '.join(ep_lines[:30])
    content_stems = _extract_stems(content_sample, min_len=5)
    # Title stems get higher weight
    all_stems = title_stems | content_stems

    if not all_stems:
        return {}

    # Determine which sūtra sections to search
    filters = set()
    if sutra_filter == 'jijnasa':
        filters = {'sutra1', 'adhyasa'}
    else:
        filters = {'sutra2'}

    # Score each sub-section
    scored = []
    for ss in subsections:
        if ss['sutra'] not in filters:
            continue
        search_text = ss['heading'] + ' ' + ss['text'][:500]
        score = 0
        matched_stems = set()
        for stem in all_stems:
            if stem in search_text:
                # Title stem = 3 points, content stem = 1 point
                weight = 3 if stem in title_stems else 1
                score += weight
                matched_stems.add(stem)
        # Require at least 2 distinct stem matches to avoid noise
        if len(matched_stems) >= 2 and score >= 4:
            scored.append((score, len(matched_stems), ss))

    # Sort by score descending
    scored.sort(key=lambda x: (-x[0], -x[1]))

    # Group by vyākhyā, take top N per vyākhyā
    result = {}
    counts = {}
    for _, _, ss in scored:
        vy_name = ss['vyakhya']
        counts.setdefault(vy_name, 0)
        if counts[vy_name] >= max_per_vyakhya:
            continue
        result.setdefault(vy_name, [])
        result[vy_name].append({
            'heading': ss['heading'],
            'text': ss['text'],
        })
        counts[vy_name] += 1

    return result


# ═══════════════════════════════════════════════════════════════════════════
# MODULE 2: DOCX FORMATTING
# ═══════════════════════════════════════════════════════════════════════════


def setup_styles(doc):
    """Set up document styles — Cambria throughout."""
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Cambria"
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2

    title = styles["Title"]
    title.font.name = "Cambria"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = DARK_BLUE
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Cambria"
    subtitle.font.size = Pt(13)
    subtitle.font.italic = True
    subtitle.font.color.rgb = MED_GRAY
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(20)

    h1 = styles["Heading 1"]
    h1.font.name = "Cambria"
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = DARK_RED
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)

    h2 = styles["Heading 2"]
    h2.font.name = "Cambria"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = DARK_BLUE
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)

    h3 = styles["Heading 3"]
    h3.font.name = "Cambria"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = DARK_GRAY
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)


def setup_page(doc):
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.0)


def add_rich_paragraph(doc, text, font_size=None, color=None, italic=False):
    """Add paragraph with inline **bold** support."""
    text = text.strip()
    if not text:
        return
    para = doc.add_paragraph()
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        if i % 2 == 1:
            run.bold = True
        if font_size:
            run.font.size = font_size
        if color:
            run.font.color.rgb = color
        if italic:
            run.italic = True
    return para


def add_section_label(doc, text, color=DARK_GRAY):
    """Add a centered section label (──── text ────)."""
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sep.add_run(f"──── {text} ────")
    run.font.size = Pt(11)
    run.font.color.rgb = color
    run.bold = True
    return sep


def add_source_text(doc, lines):
    """Add original NS Sanskrit text with distinct formatting."""
    add_section_label(doc, "न्यायसुधामूलम्", BROWN)

    for line in lines:
        s = line.strip()
        if not s:
            continue
        # Table lines (AV, Samiksha)
        if s.startswith('|'):
            text = s.strip('| ').replace('|', ' — ')
            if ':---' in text:
                continue
            # Check if this is an AV or Samiksha table marker
            if 'अनुव्याख्यानम्' in text or 'AV' in text:
                add_section_label(doc, "अनुव्याख्यानम्", BROWN)
                # Extract content after the label
                content = re.sub(r'.*?अनुव्याख्यानम्.*?:\s*', '', text).strip()
                content = re.sub(r'.*?AV.*?:\s*', '', content).strip()
                if content and content != '—':
                    para = doc.add_paragraph(content)
                    for run in para.runs:
                        run.font.size = Pt(11)
                        run.font.italic = True
                continue
            if 'समीक्षा' in text or 'Samīkṣā' in text:
                # Skip existing blank samiksha slots from source
                continue
            if text.strip() == '—' or text.strip() == '&nbsp;':
                continue
            para = doc.add_paragraph(text)
            for run in para.runs:
                run.font.size = Pt(11)
            continue

        # Regular Sanskrit text
        para = doc.add_paragraph(s)
        for run in para.runs:
            run.font.size = Pt(11)
            run.font.color.rgb = DARK_GRAY


def add_review_sections(doc, review_ep):
    """Add the Sanskrit analytical review sections."""
    if not review_ep:
        return

    for sec in review_ep.get('sections', []):
        doc.add_heading(sec['name'], level=3)
        paras = sec['text'].split('\n\n')
        for para_text in paras:
            para_text = para_text.strip()
            if not para_text:
                continue
            para_text = ' '.join(para_text.split('\n'))
            add_rich_paragraph(doc, para_text)


def add_dn_counter(doc, dn_ep, dn_num):
    """Add DN counter-refutation block."""
    add_section_label(doc, f"ध्वान्तनिरासोक्तखण्डनम् (प्रकरणम् {dn_num})", DARK_GREEN)

    # Title
    p = doc.add_paragraph()
    r = p.add_run(f"प्रकरणम् {dn_num}: {dn_ep['title']}")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = DARK_GREEN

    # Vishaya (Sanskrit only — skip vyakhya which is English)
    if dn_ep.get('vishaya'):
        doc.add_heading("विषयः", level=3)
        add_rich_paragraph(doc, dn_ep['vishaya'])

    # Purvapaksha
    if dn_ep.get('purvapaksha'):
        doc.add_heading("पूर्वपक्षः (न्यायसुधोक्तम्)", level=3)
        add_rich_paragraph(doc, dn_ep['purvapaksha'])

    # Khandana
    if dn_ep.get('khandana'):
        doc.add_heading("खण्डनम्", level=3)
        add_rich_paragraph(doc, dn_ep['khandana'])

    # Siddhanta
    if dn_ep.get('siddhanta'):
        doc.add_heading("सिद्धान्तः", level=3)
        add_rich_paragraph(doc, dn_ep['siddhanta'])


def add_vyakhya_citations(doc, vyakhya_matches):
    """Add matched vyākhyā excerpts as Advaita refutation references.

    vyakhya_matches: {vyakhya_name: [{'heading': ..., 'text': ...}, ...]}
    """
    if not vyakhya_matches:
        return

    add_section_label(doc, "अद्वैतव्याख्यानोद्धरणम्", DEEP_PURPLE)

    for vy_name, passages in vyakhya_matches.items():
        # Vyākhyā name as sub-heading
        p = doc.add_paragraph()
        r = p.add_run(vy_name)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = DEEP_PURPLE

        for passage in passages:
            # Section reference
            ref_p = doc.add_paragraph()
            ref_r = ref_p.add_run(f"「{passage['heading']}」")
            ref_r.bold = True
            ref_r.font.size = Pt(10)
            ref_r.font.color.rgb = DARK_GRAY

            # Content excerpt (first 300 chars to keep concise)
            text = passage['text'].strip()
            # Join lines, take first ~300 chars
            text = ' '.join(text.split('\n')).strip()
            if len(text) > 300:
                # Cut at last space before 300
                cut = text[:300].rfind(' ')
                if cut > 200:
                    text = text[:cut] + ' ...'
                else:
                    text = text[:300] + ' ...'

            para = doc.add_paragraph(text)
            for run in para.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = DARK_GRAY
                run.font.italic = True


def add_samiiksha_slot(doc, ep_num, part_label=""):
    """Add blank Samīkṣā reviewer notes area."""
    # Separator
    doc.add_paragraph()
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sep.add_run(f"── समीक्षा (प्रकरणम् {ep_num}) ──")
    r.font.size = Pt(11)
    r.font.color.rgb = SAFFRON
    r.bold = True

    # Field labels
    fields = [
        "मुख्यदोषः: _______________________________________________",
        "अद्वैतप्रत्युत्तरसूत्रम्: ___________________________________",
        "टिप्पणी:",
    ]
    for f in fields:
        p = doc.add_paragraph()
        run = p.add_run(f)
        run.font.size = Pt(10)
        run.font.color.rgb = LIGHT_GRAY
        run.italic = True

    # Blank lines for notes
    for _ in range(6):
        p = doc.add_paragraph()
        run = p.add_run("_" * 70)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)

    # End separator
    sep2 = doc.add_paragraph()
    sep2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sep2.add_run("── ── ──")
    r2.font.size = Pt(10)
    r2.font.color.rgb = SAFFRON
    doc.add_paragraph()


def add_bhasya_block(doc, lines):
    """Add bhāṣya or vyākhyā raw text to the document."""
    for line in lines:
        s = line.strip()
        if not s:
            continue
        if re.match(r'^---+$', s):
            continue
        # Bold markers → heading
        m = SECTION_BOLD_RE.match(s)
        if m:
            doc.add_heading(m.group(1), level=3)
            continue
        add_rich_paragraph(doc, s)


# ═══════════════════════════════════════════════════════════════════════════
# MODULE 3: DOCUMENT ASSEMBLY
# ═══════════════════════════════════════════════════════════════════════════


def build_front_matter(doc):
    """Title page and project info."""
    doc.add_paragraph("न्यायसुधा — समग्रसमीक्षा", style="Title")
    doc.add_paragraph(
        "ब्रह्मसूत्र १.१.१–१.१.२ विषये न्यायसुधायाः "
        "समीक्षात्मकविश्लेषणम्",
        style="Subtitle"
    )

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = info.add_run(
        "जिज्ञासाधिकरण-तुलनात्मक-विश्लेषण-प्रकल्पः\n"
        "अन्वेषकः: एन्. कुवलयदत्तः\n"
        "मार्गदर्शकः: प्रो. एम्.एल्.एन्. भट्टः "
        "(राष्ट्रियसंस्कृतविश्वविद्यालयः, तिरुपतिः)\n"
        "परामर्शदाता: डॉ. वंशीकृष्ण-घनपाठी"
    )
    r.font.size = Pt(10)
    r.font.color.rgb = MED_GRAY
    info.paragraph_format.space_after = Pt(30)

    doc.add_paragraph("─" * 60)

    # Reader guide
    doc.add_page_break()
    doc.add_heading("ग्रन्थसंरचना-मार्गदर्शिका", level=1)

    guide_items = [
        "**भाग अ** — जिज्ञासाधिकरणम् (ब्र.सू. १.१.१): "
        "न्यायसुधामूलम्, अनुव्याख्यानम्, संस्कृतसमीक्षा, ध्वान्तनिरासखण्डनम्, समीक्षा",
        "**भाग ब** — जन्माधिकरणम् (ब्र.सू. १.१.२): "
        "एवमेव संरचना",
        "**परिशिष्टम्** — न्यायसुधा-ध्वान्तनिरास-सम्बन्धसूची, ग्रन्थसूची",
    ]
    for item in guide_items:
        add_rich_paragraph(doc, item)

    # Episode structure guide
    doc.add_heading("प्रकरणसंरचना", level=2)
    structure_items = [
        "१. **न्यायसुधामूलम्** — जयतीर्थविरचितन्यायसुधायाः मूलसंस्कृतपाठः",
        "२. **अनुव्याख्यानम्** — मध्वाचार्यविरचितानुव्याख्यानस्य सम्बद्धश्लोकाः",
        "३. **पूर्वपक्षः / जयतीर्थयुक्तिः / अद्वैतप्रत्युत्तरम्** — समीक्षात्मकविश्लेषणम्",
        "४. **ध्वान्तनिरासोक्तखण्डनम्** — अद्वैतपक्षतः प्रत्युत्तरम् (यत्र सम्बद्धम्)",
        "५. **समीक्षा** — समीक्षकटिप्पणीनां रिक्तस्थानम्",
    ]
    for item in structure_items:
        add_rich_paragraph(doc, item)


def build_part_a(doc, bhasya_sections, sutramuktavali_text):
    """Part A: Advaita Foundational Texts."""
    doc.add_page_break()
    doc.add_heading(
        "भाग अ: अद्वैतमूलग्रन्थाः",
        level=1
    )

    for sec in bhasya_sections:
        doc.add_page_break()
        doc.add_heading(sec['heading'], level=2)

        # Main content (bhāṣya text)
        if sec['content']:
            add_bhasya_block(doc, sec['content'])

        # Vyākhyās
        for vy in sec['vyakhyas']:
            doc.add_heading(vy['name'], level=2)
            add_bhasya_block(doc, vy['lines'])

    # Sūtramuktāvalī
    if sutramuktavali_text.strip():
        doc.add_page_break()
        doc.add_heading("आनन्दगिरि-सूत्रमुक्तावली", level=2)
        for line in sutramuktavali_text.splitlines():
            s = line.strip()
            if not s:
                continue
            add_rich_paragraph(doc, s)


def build_episode_part(doc, part_label, part_key, ns_source_part,
                       review_part, dn_episodes, dn_map,
                       vyakhya_subsections=None):
    """Build Part A or B: episode-by-episode compilation."""
    doc.add_page_break()
    doc.add_heading(part_label, level=1)

    ns_reverse = _build_reverse_map(dn_map)
    all_nums = sorted(ns_source_part.keys())

    # Track which DN episodes have been included
    dn_included = set()
    vy_match_count = 0

    for num in all_nums:
        ns_ep = ns_source_part[num]
        review_ep = review_part.get(num)

        # Episode heading
        doc.add_heading(
            f"प्रकरणम् {num}: {ns_ep['title']}",
            level=2
        )

        # DN cross-reference note
        dn_nums = ns_reverse.get(num, [])
        if dn_nums:
            refs = ", ".join(str(n) for n in sorted(dn_nums))
            p = add_rich_paragraph(
                doc,
                f"**DN-सम्बन्धः**: ध्वान्तनिरासप्रकरणम् {refs}",
                font_size=Pt(10), color=DARK_GREEN, italic=True
            )

        # Layer 1: NS Original Sanskrit
        if ns_ep['lines']:
            add_source_text(doc, ns_ep['lines'])

        # Layer 2: Sanskrit Review
        if review_ep:
            add_section_label(doc, "समीक्षात्मकविश्लेषणम्", DARK_BLUE)
            add_review_sections(doc, review_ep)

        # Layer 3: DN Counter-refutation (first matching, not yet included)
        for dn_num in dn_nums:
            if dn_num not in dn_included and dn_num in dn_episodes:
                add_dn_counter(doc, dn_episodes[dn_num], dn_num)
                dn_included.add(dn_num)

        # Layer 4: Vyākhyā citations (Advaita refutation references)
        if vyakhya_subsections:
            sutra_key = 'jijnasa' if part_key == 'jijnasa' else 'janma'
            matches = find_vyakhya_for_episode(
                ns_ep['title'], ns_ep['lines'],
                vyakhya_subsections, sutra_key,
                max_per_vyakhya=2,
            )
            if matches:
                add_vyakhya_citations(doc, matches)
                vy_match_count += 1

        # Layer 5: Samīkṣā
        add_samiiksha_slot(doc, num, part_key)

    return dn_included, vy_match_count


def build_appendices(doc, dn_jijnasa, dn_janmadya, dn_included):
    """Appendices: mapping table, unmatched DN, bibliography."""
    doc.add_page_break()
    doc.add_heading("परिशिष्टम्", level=1)

    # A: Cross-reference table
    doc.add_heading("अ. न्यायसुधा-ध्वान्तनिरास-सम्बन्धसूची", level=2)
    doc.add_heading("जिज्ञासाधिकरणम् (BS 1.1.1)", level=3)
    for dn_num in sorted(DN_JIJNASA_MAP.keys()):
        ns_nums = DN_JIJNASA_MAP[dn_num]
        ns_str = ", ".join(str(n) for n in ns_nums)
        title = dn_jijnasa[dn_num]['title'] if dn_num in dn_jijnasa else "—"
        add_rich_paragraph(
            doc,
            f"**DN {dn_num}** ({title}) → NS: {ns_str}",
            font_size=Pt(10)
        )

    doc.add_heading("जन्माधिकरणम् (BS 1.1.2)", level=3)
    for dn_num in sorted(DN_JANMADYA_MAP.keys()):
        ns_nums = DN_JANMADYA_MAP[dn_num]
        ns_str = ", ".join(str(n) for n in ns_nums)
        title = dn_janmadya[dn_num]['title'] if dn_num in dn_janmadya else "—"
        add_rich_paragraph(
            doc,
            f"**DN {dn_num}** ({title}) → NS: {ns_str}",
            font_size=Pt(10)
        )

    # B: Unmatched DN prakāraṇas (full text)
    doc.add_page_break()
    doc.add_heading("आ. अप्रयुक्तध्वान्तनिरासप्रकरणानि", level=2)

    all_dn = set(dn_jijnasa.keys()) | set(dn_janmadya.keys())
    unmatched = sorted(all_dn - dn_included)

    if unmatched:
        for dn_num in unmatched:
            dn_ep = dn_jijnasa.get(dn_num) or dn_janmadya.get(dn_num)
            if dn_ep:
                add_dn_counter(doc, dn_ep, dn_num)
    else:
        add_rich_paragraph(doc, "सर्वाणि ध्वान्तनिरासप्रकरणानि मूलग्रन्थे सम्मिलितानि ।")

    # C: Bibliography placeholder
    doc.add_page_break()
    doc.add_heading("इ. सन्दर्भसूची", level=2)
    bib_entries = [
        "**न्यायसुधा** — श्रीजयतीर्थविरचिता, तत्त्वप्रकाशिका-व्याख्या, "
        "ब्रह्मसूत्रानुव्याख्यानोपरि",
        "**शारीरकमीमांसाभाष्यम्** — श्रीशङ्कराचार्यविरचितम्",
        "**आनन्दगिरिटीका** — शारीरकभाष्योपरि",
        "**ध्वान्तनिरासः** — गोडा श्रीसुब्रह्मण्यशास्त्रिभिः उपक्रान्तः",
        "**भामती** — श्रीवाचस्पतिमिश्रविरचिता",
        "**पञ्चपादिका** — श्रीपद्मपादाचार्यविरचिता",
        "**भाष्यरत्नप्रभा** — श्रीगोविन्दानन्दविरचिता",
        "**न्यायनिर्णयः** — श्रीआनन्दगिरिविरचितः",
        "**वैयासिकन्यायमाला** — श्रीभारतीतीर्थविरचिता",
        "**वेदान्तसूत्रमुक्तावली** — श्रीब्रह्मानन्दसरस्वतीविरचिता",
        "**अनुव्याख्यानम्** — श्रीमध्वाचार्यविरचितम्",
    ]
    for entry in bib_entries:
        add_rich_paragraph(doc, entry, font_size=Pt(11))


# ═══════════════════════════════════════════════════════════════════════════
# MODULE 4: MARKDOWN GENERATION
# ═══════════════════════════════════════════════════════════════════════════


def generate_markdown(ns_source, review, dn_jijnasa, dn_janmadya):
    """Generate the combined markdown reference."""
    md = []
    md.append("# न्यायसुधा — समग्रसमीक्षा\n")
    md.append("**ब्रह्मसूत्र १.१.१–१.१.२ विषये न्यायसुधायाः "
              "समीक्षात्मकविश्लेषणम्**\n")
    md.append("---\n")

    # Helper for episode parts
    def add_episode_part(part_label, part_key, dn_eps, dn_map):
        md.append(f"\n## {part_label}\n")
        ns_reverse = _build_reverse_map(dn_map)
        ns_part = ns_source[part_key]
        rev_part = review[part_key]

        for num in sorted(ns_part.keys()):
            ns_ep = ns_part[num]
            rev_ep = rev_part.get(num)
            md.append(f"\n### प्रकरणम् {num}: {ns_ep['title']}\n")

            dn_nums = ns_reverse.get(num, [])
            if dn_nums:
                refs = ", ".join(str(n) for n in sorted(dn_nums))
                md.append(f"*DN-सम्बन्धः: प्रकरणम् {refs}*\n")

            # NS source
            md.append("\n**न्यायसुधामूलम्**\n")
            for line in ns_ep['lines']:
                md.append(line)
            md.append("")

            # Review sections
            if rev_ep:
                for sec in rev_ep.get('sections', []):
                    md.append(f"\n**{sec['name']}**\n")
                    md.append(sec['text'])
                    md.append("")

            # DN counter
            for dn_num in dn_nums:
                if dn_num in dn_eps:
                    dn_ep = dn_eps[dn_num]
                    md.append(f"\n**ध्वान्तनिरासोक्तखण्डनम् (प्रकरणम् {dn_num}: "
                              f"{dn_ep['title']})**\n")
                    if dn_ep.get('vishaya'):
                        md.append(f"विषयः: {dn_ep['vishaya']}\n")
                    if dn_ep.get('khandana'):
                        md.append(f"खण्डनम्: {dn_ep['khandana']}\n")
                    if dn_ep.get('siddhanta'):
                        md.append(f"सिद्धान्तः: {dn_ep['siddhanta']}\n")

            # Samiksha
            md.append(f"\n**── समीक्षा (प्रकरणम् {num}) ──**\n")
            md.append("मुख्यदोषः: ___\n")
            md.append("अद्वैतप्रत्युत्तरसूत्रम्: ___\n")
            md.append("टिप्पणी:\n\n")
            md.append("---\n")

    # Part B
    add_episode_part(
        "भाग अ: जिज्ञासाधिकरणम् (BS 1.1.1)",
        'jijnasa', dn_jijnasa, DN_JIJNASA_MAP
    )

    # Part C
    add_episode_part(
        "भाग ब: जन्माधिकरणम् (BS 1.1.2)",
        'janma', dn_janmadya, DN_JANMADYA_MAP
    )

    content = '\n'.join(md)
    OUT_MD.write_text(content, encoding='utf-8')
    print(f"  Saved: {OUT_MD}")


# ═══════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════


def main():
    print("=" * 60)
    print("न्यायसुधा — समग्रसमीक्षा — Comprehensive Review Compilation")
    print("=" * 60)

    # ── Parse all sources ──────────────────────────────────────────────
    print("\n[1/6] Parsing NS source text...")
    ns_source = parse_ns_source(NS_SOURCE_FILE)
    j_count = len(ns_source['jijnasa'])
    m_count = len(ns_source['janma'])
    print(f"  Jijñāsā: {j_count} episodes, Janmādhi: {m_count} episodes")

    print("\n[2/6] Parsing NS Sanskrit review...")
    review = parse_review(NS_REVIEW_FILE)
    print(f"  Jijñāsā: {len(review['jijnasa'])} episodes, "
          f"Janmādhi: {len(review['janma'])} episodes")

    print("\n[3/6] Parsing DN episodes...")
    dn_jijnasa = parse_dn_episodes(DN_JIJNASA_FILE)
    dn_janmadya = parse_dn_episodes(DN_JANMADYA_FILE)
    print(f"  DN Jijñāsā: {len(dn_jijnasa)} prakāraṇas, "
          f"DN Janmādhi: {len(dn_janmadya)} prakāraṇas")

    print("\n[4/6] Parsing Advaita foundational texts...")
    bhasya_sections = parse_bhasya_sections(BHASYA_FILE)
    print(f"  Bhāṣya sections: {len(bhasya_sections)}")
    for sec in bhasya_sections:
        vy_count = len(sec['vyakhyas'])
        print(f"    {sec['heading']}: {len(sec['content'])} lines, "
              f"{vy_count} vyākhyās")

    sutramuktavali = read_raw_text(SUTRAMUKTAVALI_FILE)
    print(f"  Sūtramuktāvalī: {len(sutramuktavali)} chars")

    # ── Parse vyākhyā sub-sections for per-episode matching ────────────
    print("\n[5/7] Parsing vyākhyā sub-sections for keyword matching...")
    vy_subsections = parse_vyakhya_subsections(bhasya_sections)
    print(f"  {len(vy_subsections)} granular sub-sections parsed across 5 vyākhyās")

    # ── Assemble DOCX ──────────────────────────────────────────────────
    print("\n[6/7] Assembling DOCX...")
    doc = Document()
    setup_page(doc)
    setup_styles(doc)

    # Front matter
    build_front_matter(doc)

    # Part A: Jijñāsādhikaraṇam
    dn_included, vy_j = build_episode_part(
        doc,
        "भाग अ: जिज्ञासाधिकरणम् (ब्रह्मसूत्र १.१.१)",
        'jijnasa',
        ns_source['jijnasa'],
        review['jijnasa'],
        dn_jijnasa,
        DN_JIJNASA_MAP,
        vyakhya_subsections=vy_subsections,
    )

    # Part B: Janmādhikaraṇam
    dn_included_janma, vy_m = build_episode_part(
        doc,
        "भाग ब: जन्माधिकरणम् (ब्रह्मसूत्र १.१.२)",
        'janma',
        ns_source['janma'],
        review['janma'],
        dn_janmadya,
        DN_JANMADYA_MAP,
        vyakhya_subsections=vy_subsections,
    )
    dn_included |= dn_included_janma

    # Appendices
    build_appendices(doc, dn_jijnasa, dn_janmadya, dn_included)

    # Save
    doc.save(str(OUT_DOCX))
    print(f"  Saved: {OUT_DOCX}")

    # ── Generate Markdown ──────────────────────────────────────────────
    print("\n[7/7] Generating Markdown reference...")
    generate_markdown(ns_source, review, dn_jijnasa, dn_janmadya)

    # ── Summary ────────────────────────────────────────────────────────
    total_eps = j_count + m_count
    print(f"\n{'=' * 60}")
    print(f"Done. {total_eps} episodes compiled.")
    print(f"  NS episodes: {j_count} (Jijñāsā) + {m_count} (Janmādhi)")
    print(f"  DN matched: {len(dn_included)} prakāraṇas included inline")
    print(f"  Vyākhyā matched: {vy_j} Jijñāsā + {vy_m} Janmādhi episodes")
    print(f"  Output: {OUT_DOCX}")
    print(f"  Output: {OUT_MD}")
    print(f"{'=' * 60}")


if __name__ == "__main__":
    main()
